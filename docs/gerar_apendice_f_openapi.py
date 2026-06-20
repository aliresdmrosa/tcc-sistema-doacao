import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).parent
SOURCE = ROOT / "openapi-oficial.json"
OUTPUT = ROOT / "APENDICE_F_ESPECIFICACAO_APIS_CONEXAO_SOLIDARIA.docx"

BLACK = "000000"
GRAY = "5F6368"
LIGHT_GRAY = "F1F3F4"
BORDER = "DADCE0"

METHOD_ORDER = ["get", "post", "put", "patch", "delete", "options", "head", "trace"]
TAG_ORDER = ["Login", "Usuários", "Solicitacao", "Doação", "Reparo"]
TAG_TITLES = {
    "Login": "LOGIN",
    "Usuários": "USUÁRIOS",
    "Solicitacao": "SOLICITAÇÕES",
    "Doação": "DOAÇÕES",
    "Reparo": "REPAROS",
}

PUBLIC_OPERATIONS = {
    ("post", "/login"),
    ("post", "/usuarios"),
}


def load_spec():
    return json.loads(SOURCE.read_text(encoding="utf-8-sig"))


def iter_operations(spec):
    groups = {tag: [] for tag in TAG_ORDER}
    for path, path_item in spec.get("paths", {}).items():
        for method in METHOD_ORDER:
            if method not in path_item:
                continue
            operation = path_item[method]
            tag = (operation.get("tags") or ["Outros"])[0]
            groups.setdefault(tag, []).append((path, method, operation))
    ordered = [(tag, groups[tag]) for tag in TAG_ORDER if groups.get(tag)]
    ordered.extend((tag, ops) for tag, ops in groups.items() if tag not in TAG_ORDER and ops)
    return ordered


def schema_name(schema):
    ref = schema.get("$ref") if isinstance(schema, dict) else None
    return ref.rsplit("/", 1)[-1] if ref else ""


def resolve_ref(schema, components):
    name = schema_name(schema)
    if not name:
        return None, schema
    return name, components.get("schemas", {}).get(name, {})


def type_label(schema):
    if not isinstance(schema, dict) or not schema:
        return "object"
    if "$ref" in schema:
        return schema_name(schema)
    if "enum" in schema:
        base = schema.get("type", "string")
        return f"{base} enum: {', '.join(map(str, schema['enum']))}"
    if schema.get("type") == "array":
        return f"[{type_label(schema.get('items', {}))}]"
    label = schema.get("type", "object")
    if schema.get("format"):
        label += f" ({schema['format']})"
    if schema.get("minLength") is not None or schema.get("maxLength") is not None:
        minimum = schema.get("minLength", 0)
        maximum = schema.get("maxLength", "")
        label += f" {minimum} to {maximum} chars"
    return label


def schema_lines(schema, components, depth=0, seen=None):
    seen = set() if seen is None else seen
    if not isinstance(schema, dict) or not schema:
        return ["object"]

    if "$ref" in schema:
        name, target = resolve_ref(schema, components)
        if not name or name in seen or depth >= 3:
            return [f"{{recursive}} {name or 'object'}"]
        return schema_lines(target, components, depth, seen | {name})

    if "oneOf" in schema or "anyOf" in schema or "allOf" in schema:
        key = "oneOf" if "oneOf" in schema else "anyOf" if "anyOf" in schema else "allOf"
        lines = [f"{key}:"]
        for item in schema[key]:
            lines.extend("  " + line for line in schema_lines(item, components, depth + 1, seen))
        return lines

    if schema.get("type") == "array":
        item_lines = schema_lines(schema.get("items", {}), components, depth + 1, seen)
        if len(item_lines) == 1:
            return [f"[{item_lines[0]}]"]
        return ["[{"] + ["  " + line for line in item_lines] + ["}]"]

    properties = schema.get("properties")
    if not properties:
        return [type_label(schema)]

    required = set(schema.get("required", []))
    lines = ["{"]
    for prop_name, prop_schema in properties.items():
        marker = "*" if prop_name in required else ""
        if "$ref" in prop_schema or prop_schema.get("type") == "object":
            label = schema_name(prop_schema) or "object"
            nested = schema_lines(prop_schema, components, depth + 1, seen)
            if len(nested) > 1 and depth < 2:
                lines.append(f"{prop_name}{marker} {label} " + "{")
                lines.extend("  " + line for line in nested[1:-1])
                lines.append("}")
            else:
                lines.append(f"{prop_name}{marker} {label}")
        elif prop_schema.get("type") == "array":
            lines.append(f"{prop_name}{marker} {type_label(prop_schema)}")
        else:
            lines.append(f"{prop_name}{marker} {type_label(prop_schema)}")
    lines.append("}")
    return lines


def operation_security(spec, path, method, operation):
    if (method, path) in PUBLIC_OPERATIONS:
        return "Público"
    security = operation.get("security", spec.get("security", []))
    if security == []:
        return "Público"
    schemes = []
    for requirement in security:
        schemes.extend(requirement.keys())
    return ", ".join(dict.fromkeys(schemes)) or "Não informado"


def set_run(run, bold=False, size=10, color=BLACK, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:color"), BORDER)


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepNext"))


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_hyperlink(paragraph, text, anchor, bold=False, size=10, color="0563C1"):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:eastAsia"), "Arial")
    r_pr.append(r_fonts)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(size * 2))
    r_pr.append(size_node)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_ref(paragraph, bookmark, bold=False, size=10, color=BLACK):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark)
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run.append(begin)
    hyperlink.append(begin_run)

    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark} \\h "
    instr_run.append(instr)
    hyperlink.append(instr_run)

    sep_run = OxmlElement("w:r")
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run.append(sep)
    hyperlink.append(sep_run)

    result_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if bold:
        r_pr.append(OxmlElement("w:b"))
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(size * 2))
    r_pr.append(size_node)
    result_run.append(r_pr)
    result_text = OxmlElement("w:t")
    result_text.text = "1"
    result_run.append(result_text)
    hyperlink.append(result_run)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    hyperlink.append(end_run)
    paragraph._p.append(hyperlink)


def enable_field_updates(doc):
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 16, 16, 6, BLACK),
        ("Heading 2", 13, 12, 5, BLACK),
        ("Heading 3", 11, 10, 4, GRAY),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", 1)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Courier New"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.left_indent = Cm(0.45)
    code_style.paragraph_format.space_after = Pt(1)
    code_style.paragraph_format.line_spacing = 1.05


def add_intro(doc, spec):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("APÊNDICE F - ESPECIFICAÇÃO DE APPLICATION PROGRAMMING\nINTERFACES (APIs)")
    set_run(run, bold=True, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(
        "O documento a seguir refere-se às especificações das APIs baseadas no Swagger e OpenAPI "
        "do sistema Conexão Solidária. A documentação foi organizada no mesmo padrão de referência: "
        "identificação da API, índice por grupos funcionais, autenticação e detalhamento de cada "
        "endpoint com requisições, parâmetros, corpos e respostas HTTP."
    )
    set_run(run, size=10)

    doc.add_page_break()
    for text, size, bold, color in [
        ("API Reference", 20, True, BLACK),
        (spec.get("info", {}).get("title", "Sistema de Doação API"), 14, True, BLACK),
        (f"API Version: {spec.get('info', {}).get('version', '1.0')}", 10, False, GRAY),
        (spec.get("info", {}).get("description", ""), 10, False, GRAY),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run(run, bold=bold, size=size, color=color)


def add_index(doc, groups):
    doc.add_page_break()
    doc.add_heading("INDEX", level=1)
    for group_index, (tag, operations) in enumerate(groups, 1):
        group_anchor = f"group_{group_index}"
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9 if group_index > 1 else 2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT)
        add_hyperlink(p, f"{group_index}. {TAG_TITLES.get(tag, tag).upper()}", group_anchor, bold=True, size=10, color=BLACK)
        p.add_run("\t")
        add_page_ref(p, group_anchor, bold=True, size=10, color=BLACK)

        for endpoint_index, (path, method, _) in enumerate(operations, 1):
            anchor = f"endpoint_{group_index}_{endpoint_index}"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.15)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT)
            add_hyperlink(p, f"{group_index}.{endpoint_index} {method.upper()} {path}", anchor, size=10)
            p.add_run("\t")
            add_page_ref(p, anchor, size=10, color="0563C1")


def add_data_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = widths[index]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(header)
        set_run(run, bold=True, size=8)

    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            run = cell.paragraphs[0].add_run(str(value))
            set_run(run, size=8)
    set_table_borders(table)
    return table


def add_security(doc, spec):
    doc.add_page_break()
    doc.add_heading("Security and Authentication", level=1)
    doc.add_heading("SECURITY SCHEMES", level=2)
    schemes = spec.get("components", {}).get("securitySchemes", {})
    rows = []
    for name, scheme in schemes.items():
        scheme_type = ", ".join(filter(None, [scheme.get("type"), scheme.get("scheme"), scheme.get("bearerFormat")]))
        rows.append([name, scheme_type, scheme.get("description", "")])
    add_data_table(doc, ["KEY", "TYPE", "DESCRIPTION"], rows, [Cm(4), Cm(4.5), Cm(7.5)])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Nas operações protegidas, o token JWT deve ser enviado no cabeçalho Authorization: Bearer <token>. "
        "A configuração do backend libera publicamente POST /login, POST /usuarios, /v3/api-docs/**, "
        "/swagger-ui/** e /uploads/**. As rotas de reparo também aparecem liberadas na configuração HTTP, "
        "mas algumas operações possuem restrições por @PreAuthorize no código-fonte."
    )


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="Code Block")
        run = p.add_run(line)
        set_run(run, size=8.5, font="Courier New")


def add_parameters(doc, parameters):
    if not parameters:
        doc.add_paragraph("No request parameters")
        return
    by_location = {}
    for parameter in parameters:
        by_location.setdefault(parameter.get("in", "parameter"), []).append(parameter)

    for location, items in by_location.items():
        doc.add_heading(f"{location.upper()} PARAMETERS", level=3)
        rows = []
        for parameter in items:
            required = "*" if parameter.get("required") else ""
            rows.append(
                [
                    f"{required}{parameter.get('name', '')}",
                    type_label(parameter.get("schema", {})),
                    parameter.get("description", ""),
                ]
            )
        add_data_table(doc, ["NAME", "TYPE", "DESCRIPTION"], rows, [Cm(4), Cm(4), Cm(8)])


def add_request_body(doc, spec, request_body):
    if not request_body:
        return
    components = spec.get("components", {})
    for media_type, media in request_body.get("content", {}).items():
        doc.add_heading(f"REQUEST BODY - {media_type}", level=3)
        add_code_block(doc, schema_lines(media.get("schema", {}), components))


def add_responses(doc, spec, operation):
    components = spec.get("components", {})
    doc.add_heading("RESPONSE", level=3)
    for status, response in operation.get("responses", {}).items():
        p = doc.add_paragraph()
        run = p.add_run(f"STATUS CODE - {status}: {response.get('description', '')}")
        set_run(run, bold=True, size=9)
        content = response.get("content", {})
        for media_type, media in content.items():
            p = doc.add_paragraph()
            run = p.add_run(f"RESPONSE MODEL - {media_type}")
            set_run(run, bold=True, size=9)
            add_code_block(doc, schema_lines(media.get("schema", {}), components))


def add_operation(doc, spec, number, path, method, operation, anchor, bookmark_id):
    heading = doc.add_heading(f"{number} {method.upper()} {path}", level=2)
    add_bookmark(heading, anchor, bookmark_id)
    keep_with_next(heading)

    summary = operation.get("summary")
    description = operation.get("description")
    if summary:
        p = doc.add_paragraph()
        run = p.add_run(summary)
        set_run(run, bold=True, size=10)
    if description:
        p = doc.add_paragraph(description)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    rows = [
        ["operationId", operation.get("operationId", "não informado")],
        ["Autenticação", operation_security(spec, path, method, operation)],
    ]
    add_data_table(doc, ["ELEMENTO", "VALOR"], rows, [Cm(4), Cm(12)])

    doc.add_heading("REQUEST", level=3)
    add_parameters(doc, operation.get("parameters", []))
    add_request_body(doc, spec, operation.get("requestBody"))
    add_responses(doc, spec, operation)


def add_api(doc, spec, groups):
    doc.add_page_break()
    doc.add_heading("API", level=1)
    for group_index, (tag, operations) in enumerate(groups, 1):
        if group_index > 1:
            doc.add_page_break()
        heading = doc.add_heading(f"{group_index}. {TAG_TITLES.get(tag, tag).upper()}", level=1)
        add_bookmark(heading, f"group_{group_index}", group_index)
        tag_desc = next((item.get("description") for item in spec.get("tags", []) if item.get("name") == tag), "")
        if tag_desc:
            doc.add_paragraph(tag_desc)

        for endpoint_index, (path, method, operation) in enumerate(operations, 1):
            add_operation(
                doc,
                spec,
                f"{group_index}.{endpoint_index}",
                path,
                method,
                operation,
                f"endpoint_{group_index}_{endpoint_index}",
                1000 + group_index * 100 + endpoint_index,
            )


def build():
    spec = load_spec()
    groups = iter_operations(spec)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    configure_styles(doc)
    enable_field_updates(doc)

    add_intro(doc, spec)
    add_index(doc, groups)
    add_security(doc, spec)
    add_api(doc, spec, groups)

    props = doc.core_properties
    props.title = "Apêndice F - Especificação de APIs"
    props.subject = "Sistema Conexão Solidária"
    props.author = "Projeto Conexão Solidária"
    props.keywords = "API, Swagger, OpenAPI, TCC, Conexão Solidária"
    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"operations={sum(len(operations) for _, operations in groups)}")


if __name__ == "__main__":
    build()
