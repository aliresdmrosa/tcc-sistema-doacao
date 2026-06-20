from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Apendice_E_Especificacao_APIs.docx")

BLACK = "000000"
GRAY = "5F6368"
LIGHT_GRAY = "F1F3F4"
BORDER = "DADCE0"
METHOD_COLORS = {
    "GET": "E8F0FE",
    "POST": "E6F4EA",
    "PATCH": "FEF7E0",
    "DELETE": "FCE8E6",
}


SCHEMAS = {
    "LoginRequestDTO": [
        ("email", "string", "E-mail cadastrado do usuário."),
        ("senha", "string", "Senha do usuário."),
    ],
    "LoginDTO": [
        ("token", "string", "Token JWT para autenticação."),
        ("id", "integer (int64)", "Identificador do usuário."),
        ("email", "string", "E-mail do usuário autenticado."),
        ("perfil", "string", "Perfil: Usuario, Tecnico ou Administrador."),
    ],
    "UsuarioRequestDTO": [
        ("nome", "string", "Obrigatório. Nome completo."),
        ("cpf", "string", "CPF com 11 dígitos."),
        ("email", "string", "Obrigatório. Endereço de e-mail."),
        ("senha", "string", "Mínimo de 8 caracteres, com maiúscula, minúscula, número e símbolo."),
    ],
    "PessoaResponseDTO / UsuarioResponseDTO": [
        ("id", "integer (int64)", "Identificador do usuário."),
        ("nome", "string", "Nome completo."),
        ("cpf", "string", "CPF."),
        ("email", "string", "E-mail."),
        ("perfil", "string", "Tipo de usuário."),
        ("dataCadastro", "string (date)", "Data de cadastro."),
        ("ativo", "boolean", "Indica se o perfil está ativo."),
        ("grr", "string", "GRR, quando aplicável ao técnico."),
        ("curso", "string", "Curso, quando aplicável ao técnico."),
    ],
    "TecnicoDTO": [
        ("usuario", "UsuarioRequestDTO", "Obrigatório. Dados pessoais e de acesso."),
        ("curso", "enum Curso", "TADS, BCC, GI, HISTORIO ou BIOLOGIA."),
        ("grr", "string", "GRR com 8 dígitos."),
    ],
    "SolicitacaoRequestDTO": [
        ("equipamento", "enum Equipamento", "COMPUTADOR, NOTEBOOK, MONITOR, TECLADO ou MOUSE."),
        ("curso", "enum CursoUsuario", "Curso do solicitante."),
        ("grr", "string", "GRR com 8 dígitos."),
        ("motivo", "string", "Justificativa da solicitação."),
        ("semComputador", "boolean", "Indica ausência de computador."),
        ("ativo", "boolean", "Situação lógica da solicitação."),
    ],
    "SolicitacaoDTO": [
        ("id", "integer (int64)", "Identificador da solicitação."),
        ("grr", "string", "GRR do solicitante."),
        ("equipamento", "enum Equipamento", "Equipamento solicitado."),
        ("status", "enum Status", "Situação atual."),
        ("motivo", "string", "Justificativa."),
        ("curso", "enum CursoUsuario", "Curso do solicitante."),
        ("sem_computador", "boolean", "Indica ausência de computador."),
        ("ativo", "boolean", "Situação lógica."),
        ("dataCadastro", "string (date)", "Data de cadastro."),
        ("nome", "string", "Nome do solicitante."),
        ("cpf", "string", "CPF do solicitante."),
    ],
    "DoacaoRequestDTO (multipart/form-data)": [
        ("equipamento", "enum Equipamento", "COMPUTADOR, NOTEBOOK, MONITOR, TECLADO ou MOUSE."),
        ("quantidade", "integer", "Quantidade de itens."),
        ("descricao", "string", "Descrição da doação."),
        ("conservacao", "enum Conservacao", "NOVO, USADO ou REPARO."),
        ("imagens", "array<binary>", "Arquivos de imagem; limite global de 10 MB."),
    ],
    "DoacaoResponseDTO": [
        ("id", "integer (int64)", "Identificador da doação."),
        ("equipamento", "enum Equipamento", "Tipo de equipamento."),
        ("quantidade", "integer", "Quantidade."),
        ("descricao", "string", "Descrição."),
        ("status", "enum Status", "Situação atual."),
        ("statusConservacao", "enum Conservacao", "Estado de conservação."),
        ("dataCadastro", "string (date)", "Data de cadastro."),
        ("dataEntrega", "string (date)", "Data de entrega, quando existente."),
        ("imagens", "array<ImagemDoacao>", "Imagens relacionadas."),
    ],
    "AlterStatusDTO": [
        ("motivo", "string", "Motivo da alteração de status."),
    ],
    "ReparoResponseDTO": [
        ("id", "integer (int64)", "Identificador do reparo."),
        ("descricao", "string", "Descrição do serviço."),
        ("conclusao", "string", "Conclusão registrada."),
        ("idTecnico", "integer (int64)", "Técnico responsável."),
        ("dataInicio", "string (date)", "Data de início."),
        ("dataFim", "string (date-time)", "Data e hora de término."),
        ("idDoacao", "integer (int64)", "Doação relacionada."),
        ("equipamentoDoacao", "enum Equipamento", "Equipamento da doação."),
    ],
    "DashboardDTO": [
        ("totalUsuarios", "integer", "Total de usuários."),
        ("totalDoacoes", "integer", "Total de doações."),
        ("totalDoacoesRealizadas", "integer", "Total de doações realizadas."),
        ("doacoesAprovadas", "integer", "Doações aprovadas."),
        ("doacoesAprovadasIA", "integer", "Doações aprovadas pela análise de IA."),
        ("doacoesReprovadas", "integer", "Doações reprovadas."),
        ("doacoesReparo", "integer", "Doações em reparo."),
        ("doacoesPorMes", "array<GraficoDTO>", "Totais agrupados por mês."),
        ("doacoesPorEquipamento", "array<GraficoEquipamentoDTO>", "Totais por equipamento."),
    ],
}


GROUPS = [
    {
        "name": "LOGIN",
        "description": "Autenticação e emissão do token JWT.",
        "endpoints": [
            {
                "method": "POST", "path": "/login", "summary": "Autenticar usuário",
                "description": "Valida e-mail e senha e retorna um token JWT com os dados básicos do perfil.",
                "auth": "Público", "body": "LoginRequestDTO",
                "responses": [("200", "LoginDTO"), ("401", "Credenciais inválidas"), ("403", "Perfil desativado")],
            },
        ],
    },
    {
        "name": "USUÁRIOS",
        "description": "Cadastro, consulta, atualização e administração de perfis.",
        "endpoints": [
            {"method": "POST", "path": "/usuarios", "summary": "Cadastrar usuário", "description": "Cria um usuário comum com e-mail único.", "auth": "Público", "body": "UsuarioRequestDTO", "responses": [("201", "UsuarioResponseDTO"), ("409", "E-mail já cadastrado"), ("500", "Erro interno")]},
            {"method": "POST", "path": "/usuarios/tecnico", "summary": "Cadastrar técnico", "description": "Cria um perfil técnico com curso e GRR.", "auth": "JWT: ADMINISTRADOR", "body": "TecnicoDTO", "responses": [("201", "Tecnico"), ("403", "Acesso negado"), ("409", "Conflito no cadastro")]},
            {"method": "GET", "path": "/usuarios", "summary": "Listar usuários", "description": "Retorna todos os perfis cadastrados.", "auth": "JWT", "responses": [("200", "array<PessoaResponseDTO>"), ("500", "Erro interno")]},
            {"method": "GET", "path": "/usuarios/{id}", "summary": "Consultar usuário por ID", "description": "Retorna os dados do perfil informado.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador do usuário.")], "responses": [("200", "PessoaResponseDTO"), ("400", "ID inválido"), ("404", "Usuário não encontrado")]},
            {"method": "GET", "path": "/usuarios/cpf/{cpf}", "summary": "Consultar usuário por CPF", "description": "Busca um perfil pelo CPF.", "auth": "JWT: ADMINISTRADOR", "params": [("cpf", "path", "string", "CPF do usuário.")], "responses": [("200", "PessoaResponseDTO"), ("400", "CPF vazio"), ("404", "Usuário não encontrado")]},
            {"method": "PATCH", "path": "/usuarios/{id}", "summary": "Atualizar usuário", "description": "Atualiza os dados de um perfil existente.", "auth": "JWT: ADMINISTRADOR ou titular do perfil", "params": [("id", "path", "integer (int64)", "Identificador do usuário.")], "body": "Pessoa", "responses": [("200", "PessoaResponseDTO"), ("400", "Requisição inválida"), ("404", "Usuário não encontrado")]},
            {"method": "PATCH", "path": "/usuarios/{id}/desativar", "summary": "Desativar perfil", "description": "Desativa o perfil sem excluir seus registros.", "auth": "JWT: ADMINISTRADOR ou titular do perfil", "params": [("id", "path", "integer (int64)", "Identificador do usuário.")], "responses": [("200", "PessoaResponseDTO"), ("404", "Usuário não encontrado")]},
            {"method": "PATCH", "path": "/usuarios/{id}/reativar", "summary": "Reativar perfil", "description": "Reativa um perfil previamente desativado.", "auth": "JWT: ADMINISTRADOR", "params": [("id", "path", "integer (int64)", "Identificador do usuário.")], "responses": [("200", "PessoaResponseDTO"), ("404", "Usuário não encontrado")]},
            {"method": "DELETE", "path": "/usuarios/{id}", "summary": "Excluir usuário", "description": "Remove o usuário indicado.", "auth": "JWT: ADMINISTRADOR", "params": [("id", "path", "integer (int64)", "Identificador do usuário.")], "responses": [("204", "Excluído sem conteúdo"), ("404", "Usuário não encontrado"), ("500", "Erro interno")]},
        ],
    },
    {
        "name": "SOLICITAÇÕES",
        "description": "Registro e gestão das solicitações de equipamentos.",
        "endpoints": [
            {"method": "POST", "path": "/solicitacao", "summary": "Cadastrar solicitação", "description": "Cria uma solicitação vinculada ao usuário autenticado.", "auth": "JWT", "body": "SolicitacaoRequestDTO", "responses": [("200", "Solicitacao"), ("400", "Requisição inválida"), ("403", "Não autenticado")]},
            {"method": "GET", "path": "/solicitacao", "summary": "Listar solicitações", "description": "Retorna todas as solicitações do sistema.", "auth": "JWT", "responses": [("200", "array<SolicitacaoDTO>"), ("500", "Erro interno")]},
            {"method": "GET", "path": "/solicitacao/{id}", "summary": "Consultar solicitação por ID", "description": "Retorna a solicitação especificada.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador da solicitação.")], "responses": [("200", "SolicitacaoDTO"), ("404", "Solicitação não encontrada")]},
            {"method": "GET", "path": "/solicitacao/usuario", "summary": "Listar solicitações do usuário", "description": "Retorna as solicitações do usuário autenticado.", "auth": "JWT", "responses": [("200", "array<Solicitacao>"), ("403", "Não autenticado"), ("404", "Nenhum registro encontrado")]},
            {"method": "GET", "path": "/solicitacao/usuario/{id}", "summary": "Listar solicitações por usuário", "description": "Retorna as solicitações do usuário informado.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador do usuário.")], "responses": [("200", "array<Solicitacao>"), ("500", "Erro interno")]},
            {"method": "PATCH", "path": "/solicitacao/{id}", "summary": "Atualizar solicitação", "description": "Atualiza os dados da solicitação.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador da solicitação.")], "body": "SolicitacaoRequestDTO", "responses": [("200", "Solicitacao"), ("404", "Solicitação não encontrada")]},
            {"method": "PATCH", "path": "/solicitacao/aprovar/{id}", "summary": "Aprovar solicitação", "description": "Marca a solicitação como aprovada.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador da solicitação.")], "responses": [("200", "Sem corpo"), ("403", "Não autenticado")]},
            {"method": "PATCH", "path": "/solicitacao/reprovar/{id}", "summary": "Reprovar solicitação", "description": "Marca a solicitação como reprovada.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador da solicitação.")], "responses": [("200", "Sem corpo"), ("403", "Não autenticado")]},
            {"method": "PATCH", "path": "/solicitacao/pendente/{id}", "summary": "Reabrir análise", "description": "Retorna a solicitação ao estado pendente.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador da solicitação.")], "responses": [("200", "Sem corpo"), ("403", "Não autenticado")]},
            {"method": "PATCH", "path": "/solicitacao/{solicitacaoId}/selecionar-doacao", "summary": "Selecionar doação", "description": "Associa uma doação aprovada à solicitação.", "auth": "JWT", "params": [("solicitacaoId", "path", "integer (int64)", "Identificador da solicitação.")], "body": "integer (int64): ID da doação", "responses": [("200", "Solicitacao"), ("403", "Não autenticado")]},
            {"method": "DELETE", "path": "/solicitacao/{id}", "summary": "Excluir solicitação", "description": "Exclui a solicitação indicada.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador da solicitação.")], "responses": [("204", "Excluída sem conteúdo"), ("403", "Não autenticado")]},
        ],
    },
    {
        "name": "DOAÇÕES",
        "description": "Cadastro, consulta, avaliação e movimentação das doações.",
        "endpoints": [
            {"method": "GET", "path": "/doacao", "summary": "Listar doações", "description": "Retorna as doações cadastradas com dados resumidos do doador.", "auth": "JWT", "responses": [("200", "array<DoacaoResponseUserDTO>"), ("500", "Erro interno")]},
            {"method": "GET", "path": "/doacao/{id}", "summary": "Consultar doação por ID", "description": "Retorna os dados detalhados da doação e do doador.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador da doação.")], "responses": [("200", "DoacaoTDTO"), ("404", "Doação não encontrada")]},
            {"method": "GET", "path": "/doacao/usuario", "summary": "Listar doações do usuário", "description": "Retorna as doações do usuário autenticado.", "auth": "JWT", "responses": [("200", "array<DoacaoResponseDTO>")]},
            {"method": "GET", "path": "/doacao/usuario/{id}", "summary": "Listar doações por usuário", "description": "Retorna as doações do usuário informado.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador do usuário.")], "responses": [("200", "array<DoacaoResponseDTO>")]},
            {"method": "GET", "path": "/doacao/tipo/{equipamento}", "summary": "Filtrar por equipamento", "description": "Lista doações do tipo de equipamento informado.", "auth": "JWT", "params": [("equipamento", "path", "enum Equipamento", "Tipo do equipamento.")], "responses": [("200", "array<Doacao>"), ("500", "Erro interno")]},
            {"method": "GET", "path": "/doacao/status/{status}", "summary": "Filtrar por status", "description": "Lista doações com o status informado.", "auth": "JWT", "params": [("status", "path", "enum Status", "Status da doação.")], "responses": [("200", "array<DoacaoResponseDTO>"), ("500", "Erro interno")]},
            {"method": "GET", "path": "/doacao/avaliacao/{id}", "summary": "Consultar avaliação por IA", "description": "Retorna o histórico da avaliação automatizada da doação.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador da doação.")], "responses": [("200", "HistoricoDoacao"), ("404", "Doação não encontrada")]},
            {"method": "GET", "path": "/doacao/dashboard", "summary": "Obter dados do dashboard", "description": "Retorna indicadores e séries para cartões e gráficos.", "auth": "JWT", "responses": [("200", "DashboardDTO"), ("403", "Acesso negado"), ("500", "Erro interno")]},
            {"method": "GET", "path": "/doacao/aprovada", "summary": "Listar doações aprovadas", "description": "Lista doações disponíveis para associação a solicitações.", "auth": "JWT", "responses": [("200", "array<Doacao>"), ("500", "Erro interno")]},
            {"method": "GET", "path": "/doacao/tecnico", "summary": "Listar doações para avaliação técnica", "description": "Lista doações pendentes ou em reparo.", "auth": "JWT", "responses": [("200", "array<DoacaoReverDTO>"), ("500", "Erro interno")]},
            {"method": "POST", "path": "/doacao", "summary": "Cadastrar doação", "description": "Cria uma doação e recebe imagens por formulário multipart.", "auth": "JWT", "body": "DoacaoRequestDTO (multipart/form-data)", "responses": [("201", "DoacaoResponseDTO"), ("400", "Requisição inválida"), ("415", "Formato de arquivo inválido")]},
            {"method": "PATCH", "path": "/doacao/{id}", "summary": "Atualizar doação", "description": "Atualiza parcialmente os dados e imagens enviados.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador da doação.")], "body": "DoacaoRequestDTO (multipart/form-data)", "responses": [("200", "DoacaoResponseDTO"), ("404", "Arquivo ou doação não encontrado"), ("415", "Formato inválido")]},
            {"method": "DELETE", "path": "/doacao/{id}", "summary": "Excluir doação", "description": "Exclui a doação indicada.", "auth": "JWT", "params": [("id", "path", "integer (int64)", "Identificador da doação.")], "responses": [("204", "Excluída sem conteúdo"), ("404", "Doação não encontrada")]},
            *[
                {
                    "method": "PATCH", "path": f"/doacao/{route}/{{id}}", "summary": summary,
                    "description": description, "auth": "JWT",
                    "params": [("id", "path", "integer (int64)", "Identificador da doação.")],
                    "body": "AlterStatusDTO",
                    "responses": [("200", response), ("404", "Doação não encontrada"), ("500", "Erro interno")],
                }
                for route, summary, description, response in [
                    ("aprovar", "Aprovar doação", "Altera o status para APROVADO.", "Sem corpo"),
                    ("aprovar-reparo", "Aprovar para reparo", "Altera o status para APROVADO_REPARO.", "Sem corpo"),
                    ("entregar", "Entregar doação", "Altera o status para ENTREGUE.", "Sem corpo"),
                    ("reprovar", "Reprovar doação", "Altera o status para REPROVADO.", "Doacao"),
                    ("reparo", "Enviar para reparo", "Altera o status para REPARO.", "Doacao"),
                    ("estoque", "Enviar para estoque", "Altera o status para ESTOQUE.", "Doacao"),
                    ("pendente", "Reabrir análise", "Altera o status para PENDENTE.", "Doacao"),
                    ("aprovado-reparo", "Marcar como aprovado após reparo", "Altera o status para APROVADO_REPARO.", "Doacao"),
                    ("doado", "Marcar como doada", "Altera o status para DOADO.", "Doacao"),
                ]
            ],
        ],
    },
    {
        "name": "REPAROS",
        "description": "Acompanhamento dos serviços de reparo associados às doações.",
        "endpoints": [
            {"method": "GET", "path": "/reparo", "summary": "Listar reparos", "description": "Retorna todos os reparos cadastrados.", "auth": "Público conforme configuração atual", "responses": [("200", "array<ReparoResponseDTO>"), ("500", "Erro interno")]},
            {"method": "GET", "path": "/reparo/tecnico", "summary": "Listar reparos do técnico autenticado", "description": "Retorna os reparos atribuídos ao técnico logado.", "auth": "JWT: TECNICO", "responses": [("200", "array<ReparoResponseDTO>"), ("403", "Acesso negado")]},
            {"method": "GET", "path": "/reparo/tecnico/{id}", "summary": "Listar reparos por técnico", "description": "Retorna os reparos do técnico indicado.", "auth": "JWT: ADMINISTRADOR", "params": [("id", "path", "integer (int64)", "Identificador do técnico.")], "responses": [("200", "array<ReparoResponseDTO>"), ("403", "Acesso negado")]},
            {"method": "GET", "path": "/reparo/item/{id}", "summary": "Consultar reparo por ID", "description": "Retorna os dados de um reparo específico.", "auth": "JWT: ADMINISTRADOR ou TECNICO", "params": [("id", "path", "integer (int64)", "Identificador do reparo.")], "responses": [("200", "ReparoResponseDTO"), ("404", "Reparo não encontrado")]},
            {"method": "GET", "path": "/reparo/{id}", "summary": "Listar reparos da doação", "description": "Retorna os reparos associados à doação.", "auth": "JWT: ADMINISTRADOR ou TECNICO", "params": [("id", "path", "integer (int64)", "Identificador da doação.")], "responses": [("200", "array<ReparoResponseDTO>"), ("403", "Acesso negado")]},
            {"method": "POST", "path": "/reparo", "summary": "Cadastrar reparo", "description": "Cria um item de reparo para uma doação.", "auth": "JWT: TECNICO", "params": [("id_doacao", "query", "integer (int64)", "Identificador da doação."), ("descricao", "query", "string", "Descrição do reparo.")], "responses": [("201", "ReparoResponseDTO"), ("404", "Doação não encontrada"), ("500", "Erro interno")]},
            {"method": "PATCH", "path": "/reparo/{id}/descricao", "summary": "Atualizar descrição", "description": "Atualiza apenas a descrição do reparo.", "auth": "JWT: TECNICO", "params": [("id", "path", "integer (int64)", "Identificador do reparo.")], "body": "string: descrição", "responses": [("200", "ReparoResponseDTO"), ("404", "Reparo não encontrado")]},
            {"method": "PATCH", "path": "/reparo/concluir/{id}", "summary": "Concluir reparo", "description": "Finaliza o reparo e aprova a doação.", "auth": "JWT: TECNICO", "params": [("id", "path", "integer (int64)", "Identificador do reparo.")], "body": "string: motivo", "responses": [("200", "Sem corpo"), ("404", "Reparo não encontrado")]},
            {"method": "PATCH", "path": "/reparo/descarte/{id}", "summary": "Finalizar sem conserto", "description": "Finaliza o reparo e encaminha a doação para descarte.", "auth": "JWT: TECNICO", "params": [("id", "path", "integer (int64)", "Identificador do reparo.")], "body": "string: motivo", "responses": [("200", "Sem corpo"), ("404", "Reparo não encontrado")]},
            {"method": "PATCH", "path": "/reparo/concluir-item/{id}", "summary": "Concluir item de reparo", "description": "Marca um item de reparo como concluído.", "auth": "JWT: TECNICO", "params": [("id", "path", "integer (int64)", "Identificador do item.")], "responses": [("200", "Sem corpo"), ("404", "Item não encontrado")]},
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepNext"))


def keep_lines(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepLines"))


def set_repeat_header_font(run, bold=False, size=9, color=BLACK):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text, anchor, bold=False, size=10, color="0563C1"):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:eastAsia"), "Arial")
    r_pr.append(r_fonts)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(size * 2))
    r_pr.append(size_node)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_pageref_field(paragraph, bookmark, bold=False, size=10, color=BLACK):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark)
    hyperlink.set(qn("w:history"), "1")

    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run.append(begin)
    hyperlink.append(begin_run)

    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" PAGEREF {bookmark} \\h "
    instruction_run.append(instruction)
    hyperlink.append(instruction_run)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    hyperlink.append(separate_run)

    result_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(r_fonts)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(size * 2))
    r_pr.append(size_node)
    result_run.append(r_pr)
    result_text = OxmlElement("w:t")
    result_text.text = "1"
    result_run.append(result_text)
    hyperlink.append(result_run)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    hyperlink.append(end_run)
    paragraph._p.append(hyperlink)


def enable_field_updates(doc):
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 16, 16, 6, BLACK),
        ("Heading 2", 13, 12, 5, BLACK),
        ("Heading 3", 11, 10, 4, GRAY),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("APÊNDICE E - ESPECIFICAÇÃO DE APPLICATION PROGRAMMING\nINTERFACES (APIs)")
    set_repeat_header_font(r, bold=True, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(
        "O documento a seguir apresenta as especificações das APIs do sistema Conexão Solidária, "
        "documentadas com base no padrão OpenAPI e disponibilizadas pela interface Swagger. "
        "As operações estão organizadas por domínio funcional e descrevem os métodos HTTP, rotas, "
        "requisitos de autenticação, parâmetros, corpos de requisição e principais respostas."
    )
    set_repeat_header_font(r, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    r = p.add_run("API Reference")
    set_repeat_header_font(r, bold=True, size=20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("API do Sistema Conexão Solidária")
    set_repeat_header_font(r, bold=True, size=14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("API Version: 1.0")
    set_repeat_header_font(r, size=10, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Backend Spring Boot | OpenAPI 3 | Autenticação JWT")
    set_repeat_header_font(r, size=10, color=GRAY)


def add_index(doc):
    doc.add_page_break()
    doc.add_heading("ÍNDICE", level=1)
    for g_idx, group in enumerate(GROUPS, 1):
        group_anchor = f"group_{g_idx}"
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9 if g_idx > 1 else 2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT)
        add_internal_hyperlink(p, f"{g_idx}. {group['name']}", group_anchor, bold=True, size=10, color=BLACK)
        p.add_run("\t")
        add_pageref_field(p, group_anchor, bold=True, size=10, color=BLACK)
        for e_idx, ep in enumerate(group["endpoints"], 1):
            endpoint_anchor = f"endpoint_{g_idx}_{e_idx}"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.15)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT)
            add_internal_hyperlink(
                p,
                f"{g_idx}.{e_idx}  {ep['method']}  {ep['path']}",
                endpoint_anchor,
                size=10,
            )
            p.add_run("\t")
            add_pageref_field(p, endpoint_anchor, size=10, color="0563C1")


def add_security(doc):
    doc.add_page_break()
    doc.add_heading("SEGURANÇA E AUTENTICAÇÃO", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "A API utiliza autenticação stateless por token JWT. Nas operações protegidas, o token "
        "obtido em POST /login deve ser enviado no cabeçalho HTTP:"
    )
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Cm(0.7)
    code.paragraph_format.right_indent = Cm(0.7)
    code.paragraph_format.space_before = Pt(5)
    code.paragraph_format.space_after = Pt(8)
    run = code.add_run("Authorization: Bearer <token>")
    set_repeat_header_font(run, bold=True, size=10)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4), Cm(4), Cm(8)]
    for i, title in enumerate(["CHAVE", "TIPO", "DESCRIÇÃO"]):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(title)
        set_repeat_header_font(run, bold=True, size=8)
    row = table.add_row()
    for i, value in enumerate(["bearerAuth", "HTTP Bearer / JWT", "Autenticação usada nas rotas protegidas e nas restrições por perfil."]):
        cell = row.cells[i]
        cell.width = widths[i]
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(value)
        set_repeat_header_font(run, size=8)
    set_table_borders(table)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Rotas públicas: ")
    r.bold = True
    p.add_run("POST /login e POST /usuarios. A documentação Swagger também é pública.")
    p = doc.add_paragraph()
    r = p.add_run("Observação técnica: ")
    r.bold = True
    p.add_run(
        "a configuração HTTP atual libera /reparo/**, porém algumas operações desse grupo possuem "
        "restrições adicionais por @PreAuthorize, exigindo os perfis TECNICO ou ADMINISTRADOR."
    )


def add_schema_table(doc, name, fields):
    h = doc.add_heading(name, level=3)
    keep_with_next(h)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4.1), Cm(4.1), Cm(7.8)]
    for i, title in enumerate(["CAMPO", "TIPO", "DESCRIÇÃO"]):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(title)
        set_repeat_header_font(run, bold=True, size=8)
    set_repeat_table_header(table.rows[0])
    for field, field_type, description in fields:
        row = table.add_row()
        for i, value in enumerate([field, field_type, description]):
            cell = row.cells[i]
            cell.width = widths[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cell.paragraphs[0].add_run(value)
            set_repeat_header_font(run, bold=(i == 0), size=8)
    set_table_borders(table)


def add_schemas(doc):
    doc.add_page_break()
    doc.add_heading("MODELOS DE DADOS", level=1)
    p = doc.add_paragraph(
        "Os modelos abaixo são referenciados nas fichas dos endpoints para evitar repetição excessiva."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, fields in SCHEMAS.items():
        add_schema_table(doc, name, fields)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(4), Cm(12)]
    for label, value in rows:
        row = table.add_row()
        for i, text in enumerate([label, value]):
            cell = row.cells[i]
            cell.width = widths[i]
            set_cell_margins(cell)
            if i == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            run = cell.paragraphs[0].add_run(text)
            set_repeat_header_font(run, bold=(i == 0), size=8)
    set_table_borders(table)
    return table


def add_endpoint(doc, number, ep, anchor, bookmark_id):
    h = doc.add_heading(f"{number}  {ep['method']} {ep['path']}", level=2)
    add_bookmark(h, anchor, bookmark_id)
    keep_with_next(h)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.2)
    table.columns[1].width = Cm(13.8)
    method_cell, path_cell = table.rows[0].cells
    method_cell.width = Cm(2.2)
    path_cell.width = Cm(13.8)
    set_cell_shading(method_cell, METHOD_COLORS.get(ep["method"], LIGHT_GRAY))
    set_cell_shading(path_cell, "FFFFFF")
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=120, bottom=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    method_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = method_cell.paragraphs[0].add_run(ep["method"])
    set_repeat_header_font(run, bold=True, size=9)
    run = path_cell.paragraphs[0].add_run(ep["path"])
    set_repeat_header_font(run, bold=True, size=9)
    set_table_borders(table)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(ep["summary"])
    set_repeat_header_font(run, bold=True, size=10)
    p = doc.add_paragraph(ep["description"])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)

    rows = [("AUTENTICAÇÃO", ep["auth"])]
    if ep.get("body"):
        rows.append(("CORPO DA REQUISIÇÃO", ep["body"]))
    if ep.get("params"):
        params = "; ".join(f"{n} ({location}, {typ}): {desc}" for n, location, typ, desc in ep["params"])
        rows.append(("PARÂMETROS", params))
    responses = "; ".join(f"{code}: {desc}" for code, desc in ep["responses"])
    rows.append(("RESPOSTAS", responses))
    add_kv_table(doc, rows)


def add_endpoints(doc):
    doc.add_page_break()
    doc.add_heading("API", level=1)
    for g_idx, group in enumerate(GROUPS, 1):
        if g_idx > 1:
            doc.add_page_break()
        h = doc.add_heading(f"{g_idx}. {group['name']}", level=1)
        add_bookmark(h, f"group_{g_idx}", g_idx)
        keep_with_next(h)
        p = doc.add_paragraph(group["description"])
        p.paragraph_format.space_after = Pt(8)
        for e_idx, endpoint in enumerate(group["endpoints"], 1):
            bookmark_id = 1000 + g_idx * 100 + e_idx
            add_endpoint(
                doc,
                f"{g_idx}.{e_idx}",
                endpoint,
                f"endpoint_{g_idx}_{e_idx}",
                bookmark_id,
            )


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    configure_styles(doc)
    enable_field_updates(doc)

    add_title(doc)
    add_index(doc)
    add_security(doc)
    add_schemas(doc)
    add_endpoints(doc)

    props = doc.core_properties
    props.title = "Apêndice E - Especificação de APIs"
    props.subject = "Sistema Conexão Solidária"
    props.author = "Projeto Conexão Solidária"
    props.keywords = "API, Swagger, OpenAPI, TCC, Conexão Solidária"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
