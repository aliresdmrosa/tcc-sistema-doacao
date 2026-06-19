import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import gerar_apendice_e as base


ROOT = Path(__file__).parent
SOURCE = ROOT / "openapi-oficial.json"
FILTERED_SOURCE = ROOT / "openapi-apendice-e.json"
OUTPUT = ROOT / "APENDICE_E_OPENAPI_FINAL.docx"

METHOD_ORDER = ["get", "post", "put", "patch", "delete", "options", "head", "trace"]
TAG_ORDER = ["Login", "Usuários", "Solicitacao", "Doação", "Reparo"]
TAG_TITLES = {
    "Login": "LOGIN",
    "Usuários": "USUÁRIOS",
    "Solicitacao": "SOLICITAÇÕES",
    "Doação": "DOAÇÕES",
    "Reparo": "REPAROS",
}


def load_spec():
    spec = json.loads(SOURCE.read_text(encoding="utf-8-sig"))
    spec["paths"].pop("/usuarios/admin", None)
    spec["paths"]["/login"]["post"]["security"] = []
    spec["paths"]["/usuarios"]["post"]["security"] = []
    schemas = spec.get("components", {}).get("schemas", {})
    schemas.pop("AdministradorDTO", None)
    FILTERED_SOURCE.write_text(
        json.dumps(spec, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return spec


def iter_operations(spec):
    groups = {tag: [] for tag in TAG_ORDER}
    for path, path_item in spec["paths"].items():
        for method in METHOD_ORDER:
            if method not in path_item:
                continue
            operation = path_item[method]
            tag = (operation.get("tags") or ["Outros"])[0]
            groups.setdefault(tag, []).append((path, method, operation))
    return [(tag, groups[tag]) for tag in TAG_ORDER if groups.get(tag)]


def ref_name(schema):
    if not isinstance(schema, dict):
        return ""
    ref = schema.get("$ref")
    return ref.rsplit("/", 1)[-1] if ref else ""


def schema_label(schema):
    if not schema:
        return "não informado"
    if "$ref" in schema:
        return ref_name(schema)
    if "oneOf" in schema:
        return "oneOf<" + " | ".join(schema_label(item) for item in schema["oneOf"]) + ">"
    if "allOf" in schema:
        return "allOf<" + " + ".join(schema_label(item) for item in schema["allOf"]) + ">"
    if "anyOf" in schema:
        return "anyOf<" + " | ".join(schema_label(item) for item in schema["anyOf"]) + ">"
    schema_type = schema.get("type", "object")
    if schema_type == "array":
        return f"array<{schema_label(schema.get('items', {}))}>"
    label = schema_type
    if schema.get("format"):
        label += f" ({schema['format']})"
    if schema.get("enum"):
        label += " [" + ", ".join(map(str, schema["enum"])) + "]"
    return label


def media_schema_label(content):
    if not content:
        return "sem conteúdo"
    parts = []
    for media_type, media in content.items():
        parts.append(f"{media_type}: {schema_label(media.get('schema', {}))}")
    return "; ".join(parts)


def operation_security(spec, operation):
    security = operation.get("security", spec.get("security", []))
    if security == []:
        return "Público"
    schemes = []
    for requirement in security:
        schemes.extend(requirement.keys())
    return ", ".join(dict.fromkeys(schemes)) or "Não informado"


def add_header_row(table, labels, widths):
    row = table.rows[0]
    for index, label in enumerate(labels):
        cell = row.cells[index]
        cell.width = widths[index]
        base.set_cell_shading(cell, base.LIGHT_GRAY)
        base.set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(label)
        base.set_repeat_header_font(run, bold=True, size=8)
    base.set_repeat_table_header(row)


def add_data_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    add_header_row(table, headers, widths)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            base.set_cell_margins(cell)
            run = cell.paragraphs[0].add_run(str(value))
            base.set_repeat_header_font(run, bold=(index == 0), size=8)
    base.set_table_borders(table)
    return table


def add_title(doc, spec):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("APÊNDICE E - ESPECIFICAÇÃO DE APPLICATION PROGRAMMING\nINTERFACES (APIs)")
    base.set_repeat_header_font(run, bold=True, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(
        "O documento a seguir apresenta as especificações das APIs do sistema Conexão Solidária, "
        "geradas a partir do documento OpenAPI 3.1.0 disponibilizado pelo endpoint /v3/api-docs e "
        "visualizado pela interface Swagger UI. Para cada operação são reproduzidos os elementos "
        "definidos pela especificação: método, rota, identificação da operação, segurança, parâmetros, "
        "corpo da requisição, tipos de mídia, schemas e respostas HTTP."
    )
    base.set_repeat_header_font(run, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Nota de escopo: ")
    base.set_repeat_header_font(run, bold=True, size=9)
    run = p.add_run(
        "a operação administrativa de criação de administrador foi omitida por não integrar o fluxo utilizado pelo sistema. "
        "As operações POST /login e POST /usuarios são apresentadas como públicas, conforme a configuração "
        "de segurança da aplicação."
    )
    base.set_repeat_header_font(run, size=9)

    for text, size, bold, color in [
        ("API Reference", 20, True, base.BLACK),
        (spec["info"].get("title", "Sistema de Doação API"), 14, True, base.BLACK),
        (f"OpenAPI: {spec.get('openapi')} | API Version: {spec['info'].get('version', '')}", 10, False, base.GRAY),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        base.set_repeat_header_font(run, bold=bold, size=size, color=color)


def add_index(doc, groups):
    doc.add_page_break()
    doc.add_heading("ÍNDICE", level=1)
    for group_index, (tag, operations) in enumerate(groups, 1):
        group_anchor = f"group_{group_index}"
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9 if group_index > 1 else 2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT)
        base.add_internal_hyperlink(
            p,
            f"{group_index}. {TAG_TITLES[tag]}",
            group_anchor,
            bold=True,
            size=10,
            color=base.BLACK,
        )
        p.add_run("\t")
        base.add_pageref_field(p, group_anchor, bold=True, size=10, color=base.BLACK)
        for endpoint_index, (path, method, _) in enumerate(operations, 1):
            anchor = f"endpoint_{group_index}_{endpoint_index}"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.15)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            p.paragraph_format.tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT)
            base.add_internal_hyperlink(
                p,
                f"{group_index}.{endpoint_index}  {method.upper()}  {path}",
                anchor,
                size=10,
            )
            p.add_run("\t")
            base.add_pageref_field(p, anchor, size=10, color="0563C1")


def add_security(doc, spec):
    doc.add_page_break()
    doc.add_heading("SEGURANÇA E AUTENTICAÇÃO", level=1)
    schemes = spec.get("components", {}).get("securitySchemes", {})
    rows = []
    for name, scheme in schemes.items():
        rows.append(
            [
                name,
                scheme.get("type", ""),
                scheme.get("scheme", ""),
                scheme.get("bearerFormat", ""),
            ]
        )
    add_data_table(
        doc,
        ["CHAVE", "TIPO", "ESQUEMA", "FORMATO"],
        rows,
        [Cm(4), Cm(4), Cm(4), Cm(4)],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run(
        "Nas operações protegidas, o token deve ser enviado no cabeçalho "
        "Authorization: Bearer <token>. As fichas indicam “Público” quando a operação possui security: []."
    )


def add_operation(doc, spec, number, path, method, operation, anchor, bookmark_id):
    title = doc.add_heading(f"{number}  {method.upper()} {path}", level=2)
    base.add_bookmark(title, anchor, bookmark_id)
    base.keep_with_next(title)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    method_cell, path_cell = table.rows[0].cells
    method_cell.width = Cm(2.2)
    path_cell.width = Cm(13.8)
    base.set_cell_shading(method_cell, base.METHOD_COLORS.get(method.upper(), base.LIGHT_GRAY))
    for cell in table.rows[0].cells:
        base.set_cell_margins(cell, top=120, bottom=120)
    method_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = method_cell.paragraphs[0].add_run(method.upper())
    base.set_repeat_header_font(run, bold=True, size=9)
    run = path_cell.paragraphs[0].add_run(path)
    base.set_repeat_header_font(run, bold=True, size=9)
    base.set_table_borders(table)

    metadata = [
        ["operationId", operation.get("operationId", "não informado")],
        ["Resumo", operation.get("summary", "não informado")],
        ["Descrição", operation.get("description", "não informada")],
        ["Segurança", operation_security(spec, operation)],
    ]
    add_data_table(doc, ["ELEMENTO OPENAPI", "VALOR"], metadata, [Cm(4.2), Cm(11.8)])

    parameters = operation.get("parameters", [])
    doc.add_heading("Parameters", level=3)
    if parameters:
        parameter_rows = []
        for parameter in parameters:
            parameter_rows.append(
                [
                    parameter.get("name", ""),
                    parameter.get("in", ""),
                    "sim" if parameter.get("required") else "não",
                    schema_label(parameter.get("schema", {})),
                    parameter.get("description", ""),
                ]
            )
        add_data_table(
            doc,
            ["NAME", "IN", "REQUIRED", "SCHEMA", "DESCRIPTION"],
            parameter_rows,
            [Cm(3), Cm(2), Cm(2.2), Cm(4), Cm(4.8)],
        )
    else:
        doc.add_paragraph("No request parameters.")

    doc.add_heading("Request Body", level=3)
    request_body = operation.get("requestBody")
    if request_body:
        body_rows = []
        for media_type, media in request_body.get("content", {}).items():
            body_rows.append(
                [
                    media_type,
                    "sim" if request_body.get("required") else "não",
                    schema_label(media.get("schema", {})),
                ]
            )
        add_data_table(
            doc,
            ["CONTENT TYPE", "REQUIRED", "SCHEMA"],
            body_rows or [["não informado", "não", "não informado"]],
            [Cm(6), Cm(3), Cm(7)],
        )
    else:
        doc.add_paragraph("No request body.")

    doc.add_heading("Responses", level=3)
    response_rows = []
    for status, response in operation.get("responses", {}).items():
        response_rows.append(
            [
                status,
                response.get("description", ""),
                media_schema_label(response.get("content", {})),
            ]
        )
    add_data_table(
        doc,
        ["STATUS CODE", "DESCRIPTION", "CONTENT / SCHEMA"],
        response_rows,
        [Cm(2.7), Cm(6), Cm(7.3)],
    )


def add_operations(doc, spec, groups):
    doc.add_page_break()
    doc.add_heading("API", level=1)
    for group_index, (tag, operations) in enumerate(groups, 1):
        if group_index > 1:
            doc.add_page_break()
        heading = doc.add_heading(f"{group_index}. {TAG_TITLES[tag]}", level=1)
        base.add_bookmark(heading, f"group_{group_index}", group_index)
        for endpoint_index, (path, method, operation) in enumerate(operations, 1):
            add_operation(
                doc,
                spec,
                f"{group_index}.{endpoint_index}",
                path,
                method,
                operation,
                f"endpoint_{group_index}_{endpoint_index}",
                1000 + group_index * 100 + endpoint_index,
            )


def add_schemas(doc, spec):
    doc.add_page_break()
    doc.add_heading("COMPONENTS / SCHEMAS", level=1)
    schemas = spec.get("components", {}).get("schemas", {})
    for name, schema in schemas.items():
        heading = doc.add_heading(name, level=2)
        base.keep_with_next(heading)
        schema_type = schema.get("type", "object")
        doc.add_paragraph(f"Tipo: {schema_type}")
        properties = schema.get("properties", {})
        required = set(schema.get("required", []))
        if not properties:
            doc.add_paragraph(f"Definição: {schema_label(schema)}")
            continue
        rows = []
        for property_name, property_schema in properties.items():
            rows.append(
                [
                    property_name,
                    schema_label(property_schema),
                    "sim" if property_name in required else "não",
                    property_schema.get("description", ""),
                ]
            )
        add_data_table(
            doc,
            ["PROPERTY", "SCHEMA", "REQUIRED", "DESCRIPTION"],
            rows,
            [Cm(4.2), Cm(5.2), Cm(2.2), Cm(4.4)],
        )


def build():
    spec = load_spec()
    groups = iter_operations(spec)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    base.configure_styles(doc)
    base.enable_field_updates(doc)
    add_title(doc, spec)
    add_index(doc, groups)
    add_security(doc, spec)
    add_operations(doc, spec, groups)
    add_schemas(doc, spec)
    doc.core_properties.title = "Apêndice E - Especificação OpenAPI"
    doc.core_properties.subject = "Sistema Conexão Solidária"
    doc.core_properties.author = "Projeto Conexão Solidária"
    doc.save(OUTPUT)
    print(OUTPUT)
    print("operations", sum(len(operations) for _, operations in groups))
    print("schemas", len(spec.get("components", {}).get("schemas", {})))


if __name__ == "__main__":
    build()
